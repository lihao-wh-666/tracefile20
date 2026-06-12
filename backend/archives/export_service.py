import io
import csv
from datetime import datetime
from django.http import HttpResponse
from django.utils import timezone
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ARCHIVE_EXPORT_FIELDS = [
    ('id', 'ID'),
    ('archive_number', '案卷编号'),
    ('title', '案卷标题'),
    ('description', '案卷描述'),
    ('category_name', '所属分类'),
    ('status_display', '状态'),
    ('created_by_username', '创建人'),
    ('created_at', '创建时间'),
    ('updated_at', '更新时间'),
    ('submitted_at', '提交审核时间'),
    ('reviewed_by_username', '审核人'),
    ('reviewed_at', '审核时间'),
    ('review_comment', '审核意见'),
]


def register_chinese_font():
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        return 'STSong-Light'
    except Exception:
        return 'Helvetica'


def get_archive_field_value(archive, field_name):
    value = getattr(archive, field_name, '')
    if value is None:
        return ''
    if isinstance(value, datetime):
        return value.strftime('%Y-%m-%d %H:%M:%S')
    return str(value)


def get_special_field_value(archive, field_key):
    special_handlers = {
        'category_name': lambda a: a.category.name if a.category else '',
        'status_display': lambda a: a.get_status_display(),
        'created_by_username': lambda a: a.created_by.username if a.created_by else '',
        'reviewed_by_username': lambda a: a.reviewed_by.username if a.reviewed_by else '',
    }
    handler = special_handlers.get(field_key)
    return handler(archive) if handler else get_archive_field_value(archive, field_key)


def get_archive_export_data(archives):
    data = []
    for archive in archives:
        row = {}
        for field_key, _ in ARCHIVE_EXPORT_FIELDS:
            row[field_key] = get_special_field_value(archive, field_key)
        data.append(row)
    return data


def generate_export_filename(extension):
    return f'archives_export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.{extension}'


def create_export_response(buffer, content_type, filename):
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def export_csv(archives):
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    headers = [label for _, label in ARCHIVE_EXPORT_FIELDS]
    writer.writerow(headers)
    data = get_archive_export_data(archives)
    for row in data:
        writer.writerow([row[key] for key, _ in ARCHIVE_EXPORT_FIELDS])

    return create_export_response(
        buffer,
        'text/csv; charset=utf-8',
        generate_export_filename('csv')
    )


def export_excel(archives):
    wb = Workbook()
    ws = wb.active
    ws.title = '案卷列表'

    header_font = Font(bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    cell_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    headers = [label for _, label in ARCHIVE_EXPORT_FIELDS]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    data = get_archive_export_data(archives)
    for row_num, row_data in enumerate(data, 2):
        for col_num, (field_key, _) in enumerate(ARCHIVE_EXPORT_FIELDS, 1):
            cell = ws.cell(row=row_num, column=col_num, value=row_data[field_key])
            cell.alignment = cell_alignment
            cell.border = thin_border

    column_widths = [8, 18, 30, 40, 15, 10, 12, 18, 18, 18, 12, 18, 30]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = width

    ws.freeze_panes = 'A2'

    buffer = io.BytesIO()
    wb.save(buffer)

    return create_export_response(
        buffer,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        generate_export_filename('xlsx')
    )


def export_word(archives):
    doc = Document()

    title = doc.add_heading('案卷导出列表', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'导出时间：{timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'导出数量：{len(archives)} 条')
    doc.add_paragraph(' ')

    data = get_archive_export_data(archives)

    for idx, row_data in enumerate(data, 1):
        doc.add_heading(f'案卷 {idx}：{row_data["title"]}', level=2)

        table = doc.add_table(rows=len(ARCHIVE_EXPORT_FIELDS), cols=2)
        table.style = 'Light Grid Accent 1'

        for row_idx, (field_key, field_label) in enumerate(ARCHIVE_EXPORT_FIELDS):
            cell_label = table.cell(row_idx, 0)
            cell_value = table.cell(row_idx, 1)
            cell_label.text = field_label
            cell_value.text = row_data[field_key]

            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph(' ')

    buffer = io.BytesIO()
    doc.save(buffer)

    return create_export_response(
        buffer,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        generate_export_filename('docx')
    )


def export_pdf(archives):
    font_name = register_chinese_font()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName=font_name,
        fontSize=18,
        spaceAfter=20,
        alignment=1
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14
    )
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=12,
        textColor=colors.white
    )
    cell_style = ParagraphStyle(
        'CustomCell',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=9,
        leading=12
    )

    story = []

    story.append(Paragraph('案卷导出列表', title_style))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f'导出时间：{timezone.now().strftime("%Y-%m-%d %H:%M:%S")}', normal_style))
    story.append(Paragraph(f'导出数量：{len(archives)} 条', normal_style))
    story.append(Spacer(1, 0.5 * cm))

    table_data = []
    headers = [Paragraph(label, header_style) for _, label in ARCHIVE_EXPORT_FIELDS]
    table_data.append(headers)

    data = get_archive_export_data(archives)
    for row_data in data:
        row = []
        for field_key, _ in ARCHIVE_EXPORT_FIELDS:
            row.append(Paragraph(str(row_data[field_key]), cell_style))
        table_data.append(row)

    col_widths = [0.8 * cm, 2.2 * cm, 3.5 * cm, 4 * cm, 1.8 * cm, 1.2 * cm,
                  1.5 * cm, 2 * cm, 2 * cm, 2 * cm, 1.5 * cm, 2 * cm, 3 * cm]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
    ])
    table.setStyle(table_style)

    story.append(table)

    doc.build(story)

    return create_export_response(
        buffer,
        'application/pdf',
        generate_export_filename('pdf')
    )


EXPORT_FORMATS = {
    'pdf': export_pdf,
    'xlsx': export_excel,
    'excel': export_excel,
    'word': export_word,
    'docx': export_word,
    'csv': export_csv,
}


def export_archives(archives, export_format):
    export_format = export_format.lower().strip()
    if export_format not in EXPORT_FORMATS:
        raise ValueError(f'不支持的导出格式：{export_format}')
    export_func = EXPORT_FORMATS[export_format]
    return export_func(archives)
